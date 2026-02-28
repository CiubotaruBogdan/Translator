"""
Offline Document Translator - Server
=====================================
FastAPI backend with:
- Format-preserving DOCX translation (styles, bold, italic, fonts, tables)
- Document processing pipeline (DOCX, PDF, TXT)
- Smart chunking with paragraph preservation
- Parallel Ollama LLM translation with retry
- Configurable concurrency and num_ctx
- Server-Sent Events for real-time status
- System monitoring (CPU, RAM, GPU via Ollama)
- Connected clients tracking via WebSocket
- Job queue and history
"""

import asyncio
import copy
import json
import logging
import os
import re
import time
import uuid
import traceback
import threading
from collections import defaultdict, deque
from datetime import datetime, timedelta
from enum import Enum
from pathlib import Path
from typing import Optional

import aiohttp
import psutil
from fastapi import FastAPI, UploadFile, File, Form, HTTPException, Request, WebSocket, WebSocketDisconnect
from fastapi.responses import FileResponse, HTMLResponse, JSONResponse
from fastapi.staticfiles import StaticFiles
from starlette.responses import StreamingResponse
from starlette.middleware.cors import CORSMiddleware

import docx
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import pdfplumber
import shutil
import subprocess
import tempfile

# PDF conversion
try:
    from pdf2docx import Converter as Pdf2DocxConverter
    PDF2DOCX_AVAILABLE = True
except ImportError:
    PDF2DOCX_AVAILABLE = False

# OCR support
try:
    from PIL import Image
    import pytesseract
    TESSERACT_AVAILABLE = shutil.which('tesseract') is not None
except ImportError:
    TESSERACT_AVAILABLE = False

# ============================================================
# Configuration
# ============================================================

BASE_DIR = Path(__file__).parent
UPLOAD_DIR = BASE_DIR / "uploads"
OUTPUT_DIR = BASE_DIR / "outputs"
UPLOAD_DIR.mkdir(exist_ok=True)
OUTPUT_DIR.mkdir(exist_ok=True)

# OLLAMA_URL can be 'auto' (set by entrypoint.sh after detection) or explicit URL
_ollama_env = os.environ.get("OLLAMA_URL", "http://host.containers.internal:11434")
DEFAULT_OLLAMA_URL = _ollama_env if _ollama_env != "auto" else "http://host.containers.internal:11434"
DEFAULT_LIBRETRANSLATE_URL = os.environ.get("LIBRETRANSLATE_URL", "http://127.0.0.1:5000")
DEFAULT_MODEL = os.environ.get("OLLAMA_MODEL", "translategemma")
DEFAULT_ENGINE = os.environ.get("TRANSLATE_ENGINE", "ollama")  # 'ollama' or 'libretranslate'
MAX_CHUNK_CHARS = 1500
MAX_RETRIES = 3
RETRY_DELAY = 2
MAX_FILE_SIZE = 50 * 1024 * 1024  # 50 MB
DEFAULT_CONCURRENCY = 3
DEFAULT_NUM_CTX = 2048

LANG_EN = {
    'ro': 'Romanian', 'en': 'English', 'fr': 'French', 'de': 'German',
    'es': 'Spanish', 'it': 'Italian', 'pt': 'Portuguese', 'nl': 'Dutch',
    'pl': 'Polish', 'cs': 'Czech', 'sk': 'Slovak', 'hu': 'Hungarian',
    'bg': 'Bulgarian', 'hr': 'Croatian', 'sl': 'Slovenian', 'sr': 'Serbian',
    'sq': 'Albanian', 'mk': 'Macedonian', 'uk': 'Ukrainian', 'ru': 'Russian',
    'el': 'Greek', 'tr': 'Turkish', 'ar': 'Arabic', 'he': 'Hebrew',
    'fa': 'Persian', 'hi': 'Hindi', 'bn': 'Bengali', 'ta': 'Tamil',
    'te': 'Telugu', 'kn': 'Kannada', 'ml': 'Malayalam', 'mr': 'Marathi',
    'gu': 'Gujarati', 'ur': 'Urdu', 'ja': 'Japanese', 'ko': 'Korean',
    'zh-Hans': 'Chinese Simplified', 'zh-Hant': 'Chinese Traditional',
    'th': 'Thai', 'vi': 'Vietnamese', 'id': 'Indonesian', 'ms': 'Malay',
    'sv': 'Swedish', 'da': 'Danish', 'no': 'Norwegian', 'fi': 'Finnish',
    'et': 'Estonian', 'lt': 'Lithuanian', 'lv': 'Latvian', 'af': 'Afrikaans',
    'sw': 'Swahili', 'ca': 'Catalan', 'gl': 'Galician', 'cy': 'Welsh',
    'ga': 'Irish', 'mt': 'Maltese'
}

# ============================================================
# In-Memory Log Handler
# ============================================================

class InMemoryLogHandler(logging.Handler):
    def __init__(self, max_entries=2000):
        super().__init__()
        self.logs = deque(maxlen=max_entries)

    def emit(self, record):
        entry = {
            "time": datetime.fromtimestamp(record.created).strftime("%Y-%m-%d %H:%M:%S"),
            "level": record.levelname,
            "source": record.name,
            "message": self.format(record)
        }
        self.logs.append(entry)

log_handler = InMemoryLogHandler(max_entries=2000)
log_handler.setFormatter(logging.Formatter('%(name)s - %(message)s'))
log_handler.setLevel(logging.DEBUG)

# Attach to root logger and uvicorn loggers
logging.getLogger().addHandler(log_handler)
logging.getLogger().setLevel(logging.INFO)
for _ln in ['uvicorn', 'uvicorn.access', 'uvicorn.error', 'fastapi']:
    logging.getLogger(_ln).addHandler(log_handler)

app_logger = logging.getLogger('translator')
app_logger.setLevel(logging.DEBUG)

# ============================================================
# App Setup
# ============================================================

app = FastAPI(title="Offline Document Translator")
app.add_middleware(CORSMiddleware, allow_origins=["*"], allow_methods=["*"], allow_headers=["*"])
app.mount("/static", StaticFiles(directory=BASE_DIR / "static"), name="static")

@app.get("/logs.html")
async def serve_logs_page():
    logs_path = BASE_DIR / "static" / "logs.html"
    if logs_path.exists():
        return FileResponse(str(logs_path), media_type="text/html")
    raise HTTPException(404, "Logs page not found")

# ============================================================
# Global State
# ============================================================

class JobStatus(str, Enum):
    QUEUED = "queued"
    EXTRACTING = "extracting"
    CHUNKING = "chunking"
    TRANSLATING = "translating"
    ASSEMBLING = "assembling"
    COMPLETED = "completed"
    FAILED = "failed"
    CANCELLED = "cancelled"

class TranslationJob:
    def __init__(self, job_id, filename, source_lang, target_lang, model, ollama_url,
                 client_ip, concurrency=DEFAULT_CONCURRENCY, num_ctx=DEFAULT_NUM_CTX,
                 engine="ollama", libretranslate_url=DEFAULT_LIBRETRANSLATE_URL,
                 convert_to_pdf=False):
        self.job_id = job_id
        self.filename = filename
        self.source_lang = source_lang
        self.target_lang = target_lang
        self.model = model
        self.ollama_url = ollama_url
        self.client_ip = client_ip
        self.concurrency = concurrency
        self.num_ctx = num_ctx
        self.engine = engine  # 'ollama' or 'libretranslate'
        self.libretranslate_url = libretranslate_url
        self.convert_to_pdf = convert_to_pdf
        self.status = JobStatus.QUEUED
        self.progress = 0.0
        self.current_step = "Queued"
        self.total_chunks = 0
        self.completed_chunks = 0
        self.failed_chunks = 0
        self.retried_chunks = 0
        self.total_chars = 0
        self.translated_chars = 0
        self.error = None
        self.created_at = datetime.now().isoformat()
        self.started_at = None
        self.completed_at = None
        self.output_file = None
        self.chunks = []            # for TXT/PDF
        self.translated_chunks = [] # for TXT/PDF
        self.docx_segments = []     # for DOCX format-preserving
        self.events = []
        self.cancelled = False

    def add_event(self, etype, message, detail=""):
        self.events.append({
            "time": datetime.now().strftime("%H:%M:%S"),
            "type": etype,
            "message": message,
            "detail": detail
        })

    def to_dict(self):
        elapsed = None
        if self.started_at:
            end = self.completed_at or datetime.now().isoformat()
            try:
                t1 = datetime.fromisoformat(self.started_at)
                t2 = datetime.fromisoformat(end)
                elapsed = round((t2 - t1).total_seconds(), 1)
            except:
                pass
        return {
            "job_id": self.job_id,
            "filename": self.filename,
            "source_lang": self.source_lang,
            "target_lang": self.target_lang,
            "model": self.model,
            "client_ip": self.client_ip,
            "status": self.status.value,
            "progress": round(self.progress, 1),
            "current_step": self.current_step,
            "total_chunks": self.total_chunks,
            "completed_chunks": self.completed_chunks,
            "failed_chunks": self.failed_chunks,
            "retried_chunks": self.retried_chunks,
            "total_chars": self.total_chars,
            "translated_chars": self.translated_chars,
            "concurrency": self.concurrency,
            "num_ctx": self.num_ctx,
            "engine": self.engine,
            "error": self.error,
            "created_at": self.created_at,
            "started_at": self.started_at,
            "completed_at": self.completed_at,
            "elapsed_seconds": elapsed,
            "output_file": self.output_file,
            "events": self.events[-100:]
        }

# Stores
jobs: dict[str, TranslationJob] = {}
sse_queues: dict[str, list[asyncio.Queue]] = {}
connected_clients: dict[str, dict] = {}
server_start_time = datetime.now()

# ============================================================
# WebSocket for Dashboard
# ============================================================

dashboard_ws_clients: list[WebSocket] = []

async def broadcast_dashboard(data: dict):
    dead = []
    for ws in dashboard_ws_clients:
        try:
            await ws.send_json(data)
        except:
            dead.append(ws)
    for ws in dead:
        if ws in dashboard_ws_clients:
            dashboard_ws_clients.remove(ws)

@app.websocket("/ws/dashboard")
async def dashboard_websocket(websocket: WebSocket):
    await websocket.accept()
    dashboard_ws_clients.append(websocket)
    client_ip = websocket.client.host if websocket.client else "unknown"
    ws_id = str(uuid.uuid4())[:8]
    connected_clients[ws_id] = {
        "ip": client_ip,
        "connected_at": datetime.now().isoformat(),
        "page": "dashboard"
    }
    try:
        await websocket.send_json({"type": "init", "data": await get_dashboard_data()})
        while True:
            data = await asyncio.wait_for(websocket.receive_text(), timeout=30)
            if data == "ping":
                await websocket.send_json({"type": "pong"})
    except (WebSocketDisconnect, asyncio.TimeoutError, Exception):
        pass
    finally:
        if websocket in dashboard_ws_clients:
            dashboard_ws_clients.remove(websocket)
        connected_clients.pop(ws_id, None)

# ============================================================
# System Monitoring
# ============================================================

async def get_system_stats():
    cpu_percent = psutil.cpu_percent(interval=0)
    mem = psutil.virtual_memory()
    disk = psutil.disk_usage('/')
    gpu_info = None
    try:
        async with aiohttp.ClientSession() as session:
            async with session.get(f"{DEFAULT_OLLAMA_URL}/api/ps",
                                   timeout=aiohttp.ClientTimeout(total=3)) as resp:
                if resp.status == 200:
                    data = await resp.json()
                    if data.get("models"):
                        gpu_info = [{"model": m.get("name","unknown"),
                                     "size": m.get("size",0),
                                     "vram": m.get("size_vram",0),
                                     "expires_at": m.get("expires_at","")}
                                    for m in data["models"]]
    except:
        pass
    return {
        "cpu_percent": cpu_percent, "cpu_count": psutil.cpu_count(),
        "ram_total_gb": round(mem.total/(1024**3),1), "ram_used_gb": round(mem.used/(1024**3),1),
        "ram_percent": mem.percent,
        "disk_total_gb": round(disk.total/(1024**3),1), "disk_used_gb": round(disk.used/(1024**3),1),
        "disk_percent": round(disk.percent,1), "gpu_models": gpu_info
    }

async def get_ollama_status(url=None):
    target_url = url or DEFAULT_OLLAMA_URL
    try:
        async with aiohttp.ClientSession() as session:
            async with session.get(f"{target_url}/api/tags",
                                   timeout=aiohttp.ClientTimeout(total=5)) as resp:
                if resp.status == 200:
                    data = await resp.json()
                    models = [{"name": m.get("name",""), "size": m.get("size",0),
                               "modified_at": m.get("modified_at",""),
                               "details": m.get("details",{})}
                              for m in data.get("models",[])]
                    return {"connected": True, "models": models}
    except:
        pass
    return {"connected": False, "models": []}

async def get_libretranslate_status(url=None):
    target_url = url or DEFAULT_LIBRETRANSLATE_URL
    try:
        async with aiohttp.ClientSession() as session:
            async with session.get(f"{target_url}/languages",
                                   timeout=aiohttp.ClientTimeout(total=5)) as resp:
                if resp.status == 200:
                    langs = await resp.json()
                    return {"connected": True, "languages": len(langs), "url": target_url}
    except:
        pass
    return {"connected": False, "languages": 0, "url": target_url}

def check_libreoffice_available():
    """Check if LibreOffice headless is available."""
    try:
        result = subprocess.run(['libreoffice', '--version'],
                                capture_output=True, text=True, timeout=5)
        if result.returncode == 0:
            return {"available": True, "version": result.stdout.strip()}
    except:
        pass
    try:
        result = subprocess.run(['soffice', '--version'],
                                capture_output=True, text=True, timeout=5)
        if result.returncode == 0:
            return {"available": True, "version": result.stdout.strip()}
    except:
        pass
    return {"available": False, "version": None}

def convert_docx_to_pdf_libreoffice(docx_path: Path, output_dir: Path) -> Path:
    """Convert DOCX to PDF using LibreOffice headless."""
    try:
        cmd = ['libreoffice', '--headless', '--convert-to', 'pdf',
               '--outdir', str(output_dir), str(docx_path)]
        result = subprocess.run(cmd, capture_output=True, text=True, timeout=120)
        if result.returncode != 0:
            # Try soffice
            cmd[0] = 'soffice'
            result = subprocess.run(cmd, capture_output=True, text=True, timeout=120)
        pdf_path = output_dir / (docx_path.stem + '.pdf')
        if pdf_path.exists():
            return pdf_path
        raise Exception(f"PDF not created: {result.stderr}")
    except subprocess.TimeoutExpired:
        raise Exception("LibreOffice conversion timed out (120s)")

async def get_dashboard_data():
    sys_stats = await get_system_stats()
    ollama = await get_ollama_status()
    libretranslate = await get_libretranslate_status()
    active_jobs = [j.to_dict() for j in jobs.values()
                   if j.status in (JobStatus.QUEUED, JobStatus.EXTRACTING,
                                   JobStatus.CHUNKING, JobStatus.TRANSLATING, JobStatus.ASSEMBLING)]
    completed_jobs = sorted([j.to_dict() for j in jobs.values()
                             if j.status in (JobStatus.COMPLETED, JobStatus.FAILED, JobStatus.CANCELLED)],
                            key=lambda x: x["created_at"], reverse=True)
    uptime = (datetime.now() - server_start_time).total_seconds()
    return {
        "system": sys_stats, "ollama": ollama, "libretranslate": libretranslate,
        "connected_clients": len(connected_clients),
        "clients_detail": list(connected_clients.values()),
        "active_jobs": active_jobs, "completed_jobs": completed_jobs[:50],
        "total_jobs": len(jobs), "uptime_seconds": round(uptime)
    }

async def dashboard_updater():
    while True:
        await asyncio.sleep(3)
        if dashboard_ws_clients:
            try:
                data = await get_dashboard_data()
                await broadcast_dashboard({"type": "update", "data": data})
            except:
                pass

@app.on_event("startup")
async def startup():
    asyncio.create_task(dashboard_updater())

# ============================================================
# DOCX Format-Preserving Extraction
# ============================================================

def extract_docx_segments(filepath: Path) -> list[dict]:
    """
    Extract translatable segments from DOCX while preserving structure info.
    Each segment = one paragraph or one table cell, with metadata about location.
    Returns list of dicts: {type, index, text, char_count, ...}
    """
    doc = docx.Document(str(filepath))
    segments = []
    idx = 0

    # Body paragraphs
    for pi, para in enumerate(doc.paragraphs):
        full_text = para.text.strip()
        if full_text:
            segments.append({
                "index": idx, "type": "paragraph", "para_index": pi,
                "text": full_text, "char_count": len(full_text)
            })
            idx += 1

    # Table cells
    for ti, table in enumerate(doc.tables):
        for ri, row in enumerate(table.rows):
            for ci, cell in enumerate(row.cells):
                cell_text = cell.text.strip()
                if cell_text:
                    segments.append({
                        "index": idx, "type": "table_cell",
                        "table_index": ti, "row_index": ri, "cell_index": ci,
                        "text": cell_text, "char_count": len(cell_text)
                    })
                    idx += 1

    return segments


def apply_translations_to_docx(filepath: Path, segments: list[dict],
                                 translations: list[str], output_path: Path):
    """
    Open the original DOCX, replace text in each paragraph/cell while
    preserving formatting (runs, styles, bold, italic, fonts, etc.).
    """
    doc = docx.Document(str(filepath))

    # Build lookup: para_index -> translated text
    para_translations = {}
    cell_translations = {}

    for seg, trans in zip(segments, translations):
        if not trans or trans.startswith("[TRANSLATION ERROR"):
            continue
        if seg["type"] == "paragraph":
            para_translations[seg["para_index"]] = trans
        elif seg["type"] == "table_cell":
            key = (seg["table_index"], seg["row_index"], seg["cell_index"])
            cell_translations[key] = trans

    # Replace paragraph text preserving runs
    for pi, para in enumerate(doc.paragraphs):
        if pi in para_translations:
            _replace_paragraph_text(para, para_translations[pi])

    # Replace table cell text preserving runs
    for ti, table in enumerate(doc.tables):
        for ri, row in enumerate(table.rows):
            for ci, cell in enumerate(row.cells):
                key = (ti, ri, ci)
                if key in cell_translations:
                    # Cell can have multiple paragraphs; put all text in first, clear rest
                    if cell.paragraphs:
                        _replace_paragraph_text(cell.paragraphs[0], cell_translations[key])
                        # Clear additional paragraphs in cell
                        for extra_para in cell.paragraphs[1:]:
                            _clear_paragraph(extra_para)

    doc.save(str(output_path))


def _replace_paragraph_text(para, new_text: str):
    """
    Replace text in a paragraph while preserving the formatting of the first run.
    Strategy: put all new text in the first run, clear subsequent runs.
    This preserves the paragraph style, first run's font/bold/italic/color.
    """
    runs = para.runs
    if not runs:
        # No runs - just set text directly
        para.text = new_text
        return

    # Put translated text in first run
    runs[0].text = new_text
    # Clear all other runs
    for run in runs[1:]:
        run.text = ""


def _clear_paragraph(para):
    """Remove all text from a paragraph while keeping the XML element."""
    for run in para.runs:
        run.text = ""


# ============================================================
# Plain Text Extraction (PDF, TXT)
# ============================================================

def pdf_has_selectable_text(filepath: Path) -> bool:
    """Check if a PDF has selectable (digital) text or is a scanned image."""
    total_chars = 0
    try:
        with pdfplumber.open(str(filepath)) as pdf:
            for page in pdf.pages[:5]:  # Check first 5 pages
                text = page.extract_text()
                if text:
                    total_chars += len(text.strip())
    except:
        pass
    return total_chars > 50  # At least 50 chars = has text


def convert_pdf_to_docx(pdf_path: Path, docx_path: Path):
    """Convert PDF to DOCX using pdf2docx, preserving layout."""
    if not PDF2DOCX_AVAILABLE:
        raise Exception("pdf2docx not installed. Cannot convert PDF with format preservation.")
    cv = Pdf2DocxConverter(str(pdf_path))
    cv.convert(str(docx_path), multi_processing=False)
    cv.close()


def ocr_pdf_to_docx(pdf_path: Path, docx_path: Path, lang: str = 'eng+ron+fra'):
    """OCR a scanned PDF and create a DOCX with the extracted text."""
    if not TESSERACT_AVAILABLE:
        raise Exception(
            "Tesseract OCR not available. Cannot process scanned PDFs. "
            "Install tesseract-ocr in the container."
        )
    # Convert PDF pages to images and OCR each
    from pdf2image import convert_from_path
    images = convert_from_path(str(pdf_path), dpi=300)
    doc = docx.Document()
    for i, img in enumerate(images):
        text = pytesseract.image_to_string(img, lang=lang)
        if text.strip():
            for para in re.split(r'\n\s*\n', text):
                cleaned = para.strip()
                if cleaned:
                    doc.add_paragraph(cleaned)
        if i < len(images) - 1:
            doc.add_page_break()
    doc.save(str(docx_path))


def extract_text_from_pdf(filepath: Path) -> list[str]:
    """Fallback: extract plain text from PDF (used only if pdf2docx unavailable)."""
    paragraphs = []
    with pdfplumber.open(str(filepath)) as pdf:
        for page in pdf.pages:
            text = page.extract_text()
            if text:
                for para in re.split(r'\n\s*\n', text):
                    cleaned = para.strip().replace('\n', ' ')
                    if cleaned:
                        paragraphs.append(cleaned)
    return paragraphs

def extract_text_from_txt(filepath: Path) -> list[str]:
    text = filepath.read_text(encoding='utf-8', errors='replace')
    paragraphs = []
    for para in re.split(r'\n\s*\n', text):
        cleaned = para.strip()
        if cleaned:
            paragraphs.append(cleaned)
    return paragraphs

# ============================================================
# Smart Chunking (for plain text pipeline)
# ============================================================

def create_chunks(paragraphs: list[str], max_chars: int = MAX_CHUNK_CHARS) -> list[dict]:
    chunks = []
    current_paras = []
    current_len = 0

    for para in paragraphs:
        plen = len(para)
        if plen > max_chars:
            if current_paras:
                chunks.append({"index": len(chunks), "text": "\n\n".join(current_paras), "char_count": current_len})
                current_paras = []
                current_len = 0
            sentences = re.split(r'(?<=[.!?;])\s+', para)
            sent_buf = []
            sent_len = 0
            for sent in sentences:
                if sent_len + len(sent) > max_chars and sent_buf:
                    chunks.append({"index": len(chunks), "text": " ".join(sent_buf), "char_count": sent_len})
                    sent_buf = []
                    sent_len = 0
                sent_buf.append(sent)
                sent_len += len(sent) + 1
            if sent_buf:
                chunks.append({"index": len(chunks), "text": " ".join(sent_buf), "char_count": sent_len})
            continue
        if current_len + plen + 2 > max_chars and current_paras:
            chunks.append({"index": len(chunks), "text": "\n\n".join(current_paras), "char_count": current_len})
            current_paras = []
            current_len = 0
        current_paras.append(para)
        current_len += plen + 2

    if current_paras:
        chunks.append({"index": len(chunks), "text": "\n\n".join(current_paras), "char_count": current_len})
    return chunks

# ============================================================
# LibreTranslate Language Code Mapping
# ============================================================

# Map our lang codes to LibreTranslate codes
LIBRETRANSLATE_LANG_MAP = {
    'ro': 'ro', 'en': 'en', 'fr': 'fr', 'de': 'de', 'es': 'es', 'it': 'it',
    'pt': 'pt', 'nl': 'nl', 'pl': 'pl', 'cs': 'cs', 'sk': 'sk', 'hu': 'hu',
    'bg': 'bg', 'hr': 'hr', 'sl': 'sl', 'sr': 'sr', 'sq': 'sq', 'mk': 'mk',
    'uk': 'uk', 'ru': 'ru', 'el': 'el', 'tr': 'tr', 'ar': 'ar', 'he': 'he',
    'fa': 'fa', 'hi': 'hi', 'bn': 'bn', 'ja': 'ja', 'ko': 'ko',
    'zh-Hans': 'zh', 'zh-Hant': 'zt', 'th': 'th', 'vi': 'vi',
    'id': 'id', 'ms': 'ms', 'sv': 'sv', 'da': 'da', 'no': 'nb',
    'fi': 'fi', 'et': 'et', 'lt': 'lt', 'lv': 'lv', 'af': 'af',
    'sw': 'sw', 'ca': 'ca', 'gl': 'gl', 'cy': 'cy', 'ga': 'ga',
    'mt': 'mt', 'ta': 'ta', 'te': 'te', 'kn': 'kn', 'ml': 'ml',
    'mr': 'mr', 'gu': 'gu', 'ur': 'ur'
}

# ============================================================
# LibreTranslate Translation with Retry
# ============================================================

async def translate_chunk_libretranslate(session, text, src, tgt, lt_url):
    """Translate text using LibreTranslate API."""
    src_lt = LIBRETRANSLATE_LANG_MAP.get(src, src)
    tgt_lt = LIBRETRANSLATE_LANG_MAP.get(tgt, tgt)
    payload = {"q": text, "source": src_lt, "target": tgt_lt, "format": "text"}
    last_err = None
    for attempt in range(1, MAX_RETRIES + 1):
        try:
            async with session.post(f"{lt_url}/translate", json=payload,
                                    timeout=aiohttp.ClientTimeout(total=120)) as resp:
                if resp.status != 200:
                    err = await resp.text()
                    raise Exception(f"HTTP {resp.status}: {err[:200]}")
                data = await resp.json()
                result = data.get("translatedText", "").strip()
                if not result:
                    raise Exception("Empty response from LibreTranslate")
                return result, attempt
        except Exception as e:
            last_err = e
            if attempt < MAX_RETRIES:
                await asyncio.sleep(RETRY_DELAY * attempt)
    raise Exception(f"LibreTranslate failed after {MAX_RETRIES} attempts: {last_err}")

# ============================================================
# Ollama Translation with Retry
# ============================================================

async def translate_chunk(session, text, src, tgt, model, url,
                          num_ctx=DEFAULT_NUM_CTX, context=""):
    src_name = LANG_EN.get(src, src)
    tgt_name = LANG_EN.get(tgt, tgt)
    ctx = ""
    if context:
        ctx = f"\n\nFor context, the previous translated passage was:\n{context[-400:]}\n"
    prompt = (
        f"You are a professional {src_name} to {tgt_name} translator. "
        f"Produce only the {tgt_name} translation, no explanations or commentary."
        f"{ctx}\n"
        f"Translate the following {src_name} text into {tgt_name}:\n\n{text}"
    )
    payload = {"model": model, "prompt": prompt, "stream": False,
               "options": {"num_ctx": num_ctx}}
    last_err = None
    for attempt in range(1, MAX_RETRIES + 1):
        try:
            async with session.post(f"{url}/api/generate", json=payload,
                                    timeout=aiohttp.ClientTimeout(total=300)) as resp:
                if resp.status != 200:
                    err = await resp.text()
                    raise Exception(f"HTTP {resp.status}: {err[:200]}")
                data = await resp.json()
                result = data.get("response", "").strip()
                if not result:
                    raise Exception("Empty response from model")
                return result, attempt
        except Exception as e:
            last_err = e
            if attempt < MAX_RETRIES:
                await asyncio.sleep(RETRY_DELAY * attempt)
    raise Exception(f"Failed after {MAX_RETRIES} attempts: {last_err}")

# ============================================================
# Translation Pipeline
# ============================================================

async def run_pipeline(job: TranslationJob):
    try:
        job.started_at = datetime.now().isoformat()
        filepath = UPLOAD_DIR / f"{job.job_id}_{job.filename}"
        ext = filepath.suffix.lower()
        is_docx = ext == '.docx'
        is_pdf = ext == '.pdf'
        pdf_converted_docx = None  # Temp DOCX path if PDF was converted
        use_docx_pipeline = is_docx  # Will be set True for PDF too

        # --- EXTRACT ---
        job.status = JobStatus.EXTRACTING
        job.current_step = f"Extracting text from {job.filename}..."
        job.progress = 5
        job.add_event("info", "Extraction started", job.filename)
        emit(job)

        if is_pdf:
            # PDF: convert to DOCX first, then use format-preserving pipeline
            job.current_step = "Analyzing PDF type (digital vs scanned)..."
            emit(job)

            has_text = pdf_has_selectable_text(filepath)
            pdf_converted_docx = UPLOAD_DIR / f"{job.job_id}_converted.docx"

            if has_text:
                # Digital PDF -> convert with pdf2docx (preserves layout)
                job.current_step = "Converting PDF to DOCX (preserving layout)..."
                job.add_event("info", "PDF type: digital (selectable text)",
                              "Using pdf2docx for format preservation")
                emit(job)
                try:
                    convert_pdf_to_docx(filepath, pdf_converted_docx)
                    job.add_event("success", "PDF converted to DOCX",
                                  "Layout and formatting preserved")
                except Exception as e:
                    job.add_event("warning", f"pdf2docx conversion failed: {e}",
                                  "Falling back to plain text extraction")
                    pdf_converted_docx = None
            else:
                # Scanned PDF -> OCR then create DOCX
                job.current_step = "Scanned PDF detected. Running OCR..."
                # Map source language to Tesseract language codes
                tess_lang_map = {
                    'ro': 'ron', 'en': 'eng', 'fr': 'fra', 'de': 'deu',
                    'es': 'spa', 'it': 'ita', 'pt': 'por', 'nl': 'nld',
                    'pl': 'pol', 'ru': 'rus', 'uk': 'ukr', 'bg': 'bul',
                    'hr': 'hrv', 'cs': 'ces', 'sk': 'slk', 'hu': 'hun',
                    'el': 'ell', 'tr': 'tur', 'ar': 'ara', 'ja': 'jpn',
                    'ko': 'kor', 'zh-Hans': 'chi_sim', 'zh-Hant': 'chi_tra'
                }
                ocr_lang = tess_lang_map.get(job.source_lang, 'eng')
                # Always include eng as fallback
                if ocr_lang != 'eng':
                    ocr_lang = f"{ocr_lang}+eng"
                job.add_event("info", "PDF type: scanned (image-based)",
                              f"Running Tesseract OCR with lang={ocr_lang}")
                emit(job)
                try:
                    ocr_pdf_to_docx(filepath, pdf_converted_docx, lang=ocr_lang)
                    job.add_event("success", "OCR complete, DOCX created",
                                  "Text extracted from scanned pages")
                except Exception as e:
                    job.add_event("warning", f"OCR failed: {e}",
                                  "Falling back to plain text extraction")
                    pdf_converted_docx = None

            if pdf_converted_docx and pdf_converted_docx.exists():
                # Use the converted DOCX for format-preserving pipeline
                filepath = pdf_converted_docx
                use_docx_pipeline = True
            else:
                # Fallback: plain text extraction
                use_docx_pipeline = False

        if use_docx_pipeline:
            # Format-preserving DOCX pipeline (for DOCX and converted PDF)
            try:
                segments = extract_docx_segments(filepath)
            except Exception as e:
                raise Exception(f"DOCX extraction failed: {e}")
            if not segments:
                raise Exception("No text content found in document")

            job.docx_segments = segments
            job.total_chars = sum(s["char_count"] for s in segments)
            job.total_chunks = len(segments)
            job.translated_chunks = [""] * len(segments)
            mode_label = "PDF->DOCX format-preserving" if is_pdf else "DOCX format-preserving"
            job.add_event("success", f"Extraction complete ({mode_label})",
                          f"{len(segments)} segments, {job.total_chars:,} chars")
        elif not is_docx:
            # Plain text pipeline (TXT or PDF fallback)
            try:
                if is_pdf:
                    paragraphs = extract_text_from_pdf(filepath)
                else:
                    paragraphs = extract_text_from_txt(filepath)
            except Exception as e:
                raise Exception(f"Extraction failed: {e}")
            if not paragraphs:
                raise Exception("No text content found in document")

            job.total_chars = sum(len(p) for p in paragraphs)
            job.add_event("success", "Extraction complete (plain text)",
                          f"{len(paragraphs)} paragraphs, {job.total_chars:,} chars")

        emit(job)

        # --- CHUNK (only for non-DOCX) ---
        if use_docx_pipeline:
            # For DOCX (and converted PDF), each segment is a chunk
            chunks_to_translate = [{"index": s["index"], "text": s["text"],
                                    "char_count": s["char_count"]} for s in segments]
            job.add_event("info", "Using format-preserving mode",
                          f"{len(chunks_to_translate)} segments (1 per paragraph/cell)")
        else:
            job.status = JobStatus.CHUNKING
            job.current_step = "Splitting into translation chunks..."
            job.progress = 10
            emit(job)
            chunks_to_translate = create_chunks(paragraphs)
            job.total_chunks = len(chunks_to_translate)
            job.translated_chunks = [""] * len(chunks_to_translate)
            job.add_event("info", "Chunking complete",
                          f"{len(chunks_to_translate)} chunks")

        engine_info = f"engine={job.engine}"
        if job.engine == 'libretranslate':
            engine_info += f", url={job.libretranslate_url}"
        else:
            engine_info += f", model={job.model}, num_ctx={job.num_ctx}"
        job.add_event("info", "Pipeline config",
                      f"concurrency={job.concurrency}, {engine_info}")
        emit(job)

        # --- TRANSLATE (parallel) ---
        job.status = JobStatus.TRANSLATING
        job.progress = 12
        emit(job)

        use_libretranslate = (job.engine == 'libretranslate')

        async with aiohttp.ClientSession() as session:
            # Verify connection to translation engine
            if use_libretranslate:
                try:
                    async with session.get(f"{job.libretranslate_url}/languages",
                                           timeout=aiohttp.ClientTimeout(total=5)) as r:
                        if r.status != 200:
                            raise Exception(f"HTTP {r.status}")
                except Exception as e:
                    raise Exception(f"Cannot connect to LibreTranslate at {job.libretranslate_url}: {e}")
                job.add_event("success", "LibreTranslate connected", job.libretranslate_url)
            else:
                try:
                    async with session.get(f"{job.ollama_url}/api/tags",
                                           timeout=aiohttp.ClientTimeout(total=5)) as r:
                        if r.status != 200:
                            raise Exception(f"HTTP {r.status}")
                except Exception as e:
                    raise Exception(f"Cannot connect to Ollama at {job.ollama_url}: {e}")
                job.add_event("success", "Ollama connected", job.ollama_url)

            semaphore = asyncio.Semaphore(job.concurrency)
            lock = asyncio.Lock()

            async def translate_one(i, chunk):
                if job.cancelled:
                    return
                async with semaphore:
                    if job.cancelled:
                        return
                    n = i + 1
                    async with lock:
                        engine_label = "LibreTranslate" if use_libretranslate else "Ollama"
                        job.current_step = f"[{engine_label}] Translating segment {n}/{len(chunks_to_translate)} ({chunk['char_count']} chars)..."
                        job.progress = 12 + (78 * job.completed_chunks / len(chunks_to_translate))
                        emit(job)
                    try:
                        if use_libretranslate:
                            translated, attempts = await translate_chunk_libretranslate(
                                session, chunk["text"], job.source_lang, job.target_lang,
                                job.libretranslate_url
                            )
                        else:
                            translated, attempts = await translate_chunk(
                                session, chunk["text"], job.source_lang, job.target_lang,
                                job.model, job.ollama_url, num_ctx=job.num_ctx
                            )
                        job.translated_chunks[i] = translated
                        async with lock:
                            job.completed_chunks += 1
                            job.translated_chars += len(translated)
                            job.progress = 12 + (78 * job.completed_chunks / len(chunks_to_translate))
                            if attempts > 1:
                                job.retried_chunks += 1
                                job.add_event("warning", f"Segment {n}/{len(chunks_to_translate)} OK after {attempts} attempts",
                                              f"{chunk['char_count']}\u2192{len(translated)} chars")
                            else:
                                job.add_event("success", f"Segment {n}/{len(chunks_to_translate)} translated",
                                              f"{chunk['char_count']}\u2192{len(translated)} chars")
                            emit(job)
                    except Exception as e:
                        async with lock:
                            job.failed_chunks += 1
                            job.completed_chunks += 1
                            job.translated_chunks[i] = f"[TRANSLATION ERROR: {e}]"
                            job.progress = 12 + (78 * job.completed_chunks / len(chunks_to_translate))
                            job.add_event("error", f"Segment {n}/{len(chunks_to_translate)} FAILED", str(e))
                            emit(job)

            tasks = [translate_one(i, chunk) for i, chunk in enumerate(chunks_to_translate)]
            await asyncio.gather(*tasks)

            if job.cancelled:
                job.status = JobStatus.CANCELLED
                job.current_step = "Cancelled by user"
                job.add_event("warning", "Job cancelled by user")
                job.completed_at = datetime.now().isoformat()
                emit(job)
                return

        # --- ASSEMBLE ---
        job.status = JobStatus.ASSEMBLING
        job.current_step = "Assembling translated document..."
        job.progress = 92
        emit(job)

        if use_docx_pipeline:
            out_name = Path(job.filename).stem + f"_{job.target_lang}.docx"
            output_path = OUTPUT_DIR / f"{job.job_id}_{out_name}"
            apply_translations_to_docx(filepath, segments, job.translated_chunks, output_path)
            label = "PDF->DOCX assembled (format preserved)" if is_pdf else "DOCX assembled (format preserved)"
            job.add_event("success", label, out_name)
            # Cleanup temp converted DOCX if it was a PDF
            if pdf_converted_docx and pdf_converted_docx.exists():
                try:
                    pdf_converted_docx.unlink()
                except:
                    pass
        else:
            out_name = Path(job.filename).stem + f"_{job.target_lang}"
            if ext == '.pdf':
                out_name += '.txt'
            else:
                out_name += ext
            assemble_txt(job, out_name)
            job.add_event("success", "Document assembled", out_name)

        job.output_file = out_name

        # --- CONVERT TO PDF (optional, LibreOffice headless) ---
        if job.convert_to_pdf and out_name.endswith('.docx'):
            try:
                job.current_step = "Converting to PDF (LibreOffice)..."
                job.progress = 96
                emit(job)
                docx_output_path = OUTPUT_DIR / f"{job.job_id}_{out_name}"
                pdf_path = convert_docx_to_pdf_libreoffice(docx_output_path, OUTPUT_DIR)
                # Rename to include job_id prefix
                pdf_name = Path(job.filename).stem + f"_{job.target_lang}.pdf"
                final_pdf = OUTPUT_DIR / f"{job.job_id}_{pdf_name}"
                if pdf_path != final_pdf:
                    pdf_path.rename(final_pdf)
                # Remove DOCX, keep only PDF
                if docx_output_path.exists():
                    docx_output_path.unlink()
                job.output_file = pdf_name
                job.add_event("success", "Converted to PDF", pdf_name)
            except Exception as e:
                job.add_event("warning", "PDF conversion failed, keeping DOCX", str(e))
                app_logger.warning(f"LibreOffice conversion failed: {e}")

        # --- DONE ---
        job.status = JobStatus.COMPLETED
        job.progress = 100
        job.completed_at = datetime.now().isoformat()
        job.current_step = "Complete!"
        job.add_event("success", "Translation complete",
                      f"{job.completed_chunks}/{job.total_chunks} segments, "
                      f"{job.failed_chunks} failures, {job.retried_chunks} retries")
        emit(job)

    except Exception as e:
        job.status = JobStatus.FAILED
        job.error = str(e)
        job.current_step = f"Failed: {e}"
        job.completed_at = datetime.now().isoformat()
        job.add_event("error", "Pipeline failed", str(e))
        emit(job)

def assemble_txt(job, out_name):
    path = OUTPUT_DIR / f"{job.job_id}_{out_name}"
    with open(path, 'w', encoding='utf-8') as f:
        for translated in job.translated_chunks:
            if translated:
                f.write(translated + "\n\n")

# ============================================================
# SSE
# ============================================================

def emit(job: TranslationJob):
    jid = job.job_id
    if jid in sse_queues:
        data = job.to_dict()
        for q in sse_queues[jid]:
            try:
                q.put_nowait(data)
            except asyncio.QueueFull:
                pass
    asyncio.ensure_future(broadcast_dashboard({"type": "job_update", "data": job.to_dict()}))

# ============================================================
# API Endpoints
# ============================================================

@app.get("/", response_class=HTMLResponse)
async def index():
    return HTMLResponse((BASE_DIR / "static" / "index.html").read_text(encoding='utf-8'))

@app.get("/api/status")
async def api_status():
    return await get_dashboard_data()

@app.get("/api/ollama")
async def api_ollama(url: str = DEFAULT_OLLAMA_URL):
    return await get_ollama_status(url)

@app.get("/api/models")
async def api_models(url: str = DEFAULT_OLLAMA_URL):
    result = await get_ollama_status(url)
    return {"connected": result["connected"], "models": result["models"]}

@app.get("/api/libretranslate")
async def api_libretranslate(url: str = DEFAULT_LIBRETRANSLATE_URL):
    """Check LibreTranslate status and available languages."""
    try:
        async with aiohttp.ClientSession() as session:
            async with session.get(f"{url}/languages",
                                   timeout=aiohttp.ClientTimeout(total=5)) as resp:
                if resp.status == 200:
                    languages = await resp.json()
                    return {"connected": True, "languages": languages, "url": url}
    except:
        pass
    return {"connected": False, "languages": [], "url": url}

@app.post("/api/translate")
async def api_translate(
    request: Request,
    file: UploadFile = File(...),
    source_lang: str = Form("ro"),
    target_lang: str = Form("en"),
    model: str = Form(DEFAULT_MODEL),
    ollama_url: str = Form(DEFAULT_OLLAMA_URL),
    concurrency: int = Form(DEFAULT_CONCURRENCY),
    num_ctx: int = Form(DEFAULT_NUM_CTX),
    engine: str = Form(DEFAULT_ENGINE),
    libretranslate_url: str = Form(DEFAULT_LIBRETRANSLATE_URL),
    convert_to_pdf: str = Form("false")
):
    ext = Path(file.filename).suffix.lower()
    if ext not in ('.docx', '.pdf', '.txt', '.text', '.md'):
        raise HTTPException(400, f"Unsupported: {ext}")
    content = await file.read()
    if len(content) > MAX_FILE_SIZE:
        raise HTTPException(400, f"File too large (max {MAX_FILE_SIZE // 1024 // 1024}MB)")
    concurrency = max(1, min(concurrency, 10))
    num_ctx = max(512, min(num_ctx, 32768))
    job_id = str(uuid.uuid4())[:8]
    client_ip = request.client.host if request.client else "unknown"
    engine = engine if engine in ('ollama', 'libretranslate') else 'ollama'
    pdf_convert = convert_to_pdf.lower() in ('true', '1', 'yes')
    job = TranslationJob(job_id, file.filename, source_lang, target_lang, model, ollama_url,
                         client_ip, concurrency=concurrency, num_ctx=num_ctx,
                         engine=engine, libretranslate_url=libretranslate_url,
                         convert_to_pdf=pdf_convert)
    jobs[job_id] = job
    save_path = UPLOAD_DIR / f"{job_id}_{file.filename}"
    save_path.write_bytes(content)
    job.add_event("info", "File uploaded", f"{file.filename} ({len(content):,} bytes) from {client_ip}")
    asyncio.create_task(run_pipeline(job))
    return {"job_id": job_id, "status": "queued"}

@app.get("/api/jobs")
async def api_jobs():
    return {"jobs": [j.to_dict() for j in sorted(jobs.values(), key=lambda j: j.created_at, reverse=True)]}

@app.get("/api/jobs/{job_id}")
async def api_job(job_id: str):
    if job_id not in jobs:
        raise HTTPException(404, "Job not found")
    return jobs[job_id].to_dict()

@app.post("/api/jobs/{job_id}/cancel")
async def api_cancel(job_id: str):
    if job_id not in jobs:
        raise HTTPException(404)
    jobs[job_id].cancelled = True
    return {"ok": True}

@app.get("/api/jobs/{job_id}/download")
async def api_download(job_id: str):
    if job_id not in jobs:
        raise HTTPException(404)
    job = jobs[job_id]
    if not job.output_file:
        raise HTTPException(400, "Not ready")
    path = OUTPUT_DIR / f"{job_id}_{job.output_file}"
    if not path.exists():
        raise HTTPException(404, "File missing")
    return FileResponse(str(path), filename=job.output_file)

@app.delete("/api/jobs/{job_id}")
async def api_delete(job_id: str):
    if job_id not in jobs:
        raise HTTPException(404)
    job = jobs.pop(job_id)
    for p in [UPLOAD_DIR / f"{job_id}_{job.filename}",
              OUTPUT_DIR / f"{job_id}_{job.output_file}" if job.output_file else None]:
        if p and p.exists():
            p.unlink()
    return {"ok": True}

@app.post("/api/jobs/stop-all")
async def api_stop_all():
    stopped = 0
    for job in jobs.values():
        if job.status in (JobStatus.QUEUED, JobStatus.EXTRACTING,
                          JobStatus.CHUNKING, JobStatus.TRANSLATING, JobStatus.ASSEMBLING):
            job.cancelled = True
            stopped += 1
    app_logger.warning(f"Stop all: {stopped} jobs cancelled")
    return {"ok": True, "stopped": stopped}

# ============================================================
# Logs & Files Endpoints
# ============================================================

@app.get("/api/logs")
async def api_logs(limit: int = 500, level: str = ""):
    logs = list(log_handler.logs)
    if level:
        logs = [l for l in logs if l["level"] == level.upper()]
    return {"logs": logs[-limit:], "total": len(log_handler.logs)}

@app.get("/api/files")
async def api_files():
    files = []
    for d, category in [(UPLOAD_DIR, "upload"), (OUTPUT_DIR, "output")]:
        if d.exists():
            for f in sorted(d.iterdir(), key=lambda x: x.stat().st_mtime, reverse=True):
                if f.is_file():
                    stat = f.stat()
                    files.append({
                        "name": f.name,
                        "category": category,
                        "size": stat.st_size,
                        "modified": datetime.fromtimestamp(stat.st_mtime).isoformat(),
                        "path": str(f.relative_to(BASE_DIR))
                    })
    return {"files": files, "upload_dir": str(UPLOAD_DIR), "output_dir": str(OUTPUT_DIR)}

@app.delete("/api/files/{category}/{filename}")
async def api_delete_file(category: str, filename: str):
    if category == "upload":
        path = UPLOAD_DIR / filename
    elif category == "output":
        path = OUTPUT_DIR / filename
    else:
        raise HTTPException(400, "Invalid category")
    if not path.exists():
        raise HTTPException(404, "File not found")
    path.unlink()
    app_logger.info(f"File deleted: {category}/{filename}")
    return {"ok": True}

@app.post("/api/files/cleanup")
async def api_cleanup_files():
    deleted = 0
    for d in [UPLOAD_DIR, OUTPUT_DIR]:
        if d.exists():
            for f in d.iterdir():
                if f.is_file():
                    f.unlink()
                    deleted += 1
    app_logger.warning(f"Cleanup: {deleted} files deleted")
    return {"ok": True, "deleted": deleted}

@app.get("/api/libreoffice")
async def api_libreoffice():
    return check_libreoffice_available()

@app.get("/api/system-info")
async def api_system_info():
    import platform
    return {
        "platform": platform.platform(),
        "python": platform.python_version(),
        "cpu_count": psutil.cpu_count(),
        "ram_total_gb": round(psutil.virtual_memory().total / (1024**3), 1),
        "server_start": server_start_time.isoformat(),
        "uptime_seconds": round((datetime.now() - server_start_time).total_seconds()),
        "upload_dir": str(UPLOAD_DIR),
        "output_dir": str(OUTPUT_DIR),
        "total_jobs": len(jobs),
        "active_jobs": sum(1 for j in jobs.values() if j.status in
                          (JobStatus.QUEUED, JobStatus.EXTRACTING, JobStatus.CHUNKING,
                           JobStatus.TRANSLATING, JobStatus.ASSEMBLING)),
        "ollama_url": DEFAULT_OLLAMA_URL
    }

@app.get("/api/network-diag")
async def api_network_diagnostics():
    """Run network diagnostics to help troubleshoot Ollama connectivity."""
    import subprocess as sp
    results = {"tests": [], "ollama_url": DEFAULT_OLLAMA_URL}

    # Get container network info
    try:
        gw = sp.run(['ip', 'route'], capture_output=True, text=True, timeout=5)
        results["routes"] = gw.stdout.strip()
    except:
        results["routes"] = "unavailable"

    try:
        hostname_out = sp.run(['hostname', '-I'], capture_output=True, text=True, timeout=5)
        results["container_ip"] = hostname_out.stdout.strip()
    except:
        results["container_ip"] = "unavailable"

    # Test various host endpoints
    ollama_port = os.environ.get("OLLAMA_PORT", "11434")
    test_hosts = [
        "host.containers.internal",
        "host.docker.internal",
        "host.wsl.internal",
    ]
    # Add gateway IP
    try:
        gw_line = sp.run(['ip', 'route', 'show', 'default'], capture_output=True, text=True, timeout=5)
        gw_ip = gw_line.stdout.strip().split()[2] if gw_line.stdout.strip() else None
        if gw_ip:
            test_hosts.append(gw_ip)
            results["gateway_ip"] = gw_ip
    except:
        pass

    for host in test_hosts:
        test = {"host": host, "ping": False, "ollama": False, "resolved_ip": None}
        # Ping test
        try:
            ping = sp.run(['ping', '-c', '1', '-W', '2', host],
                          capture_output=True, text=True, timeout=5)
            test["ping"] = (ping.returncode == 0)
        except:
            pass
        # DNS resolution
        try:
            getent = sp.run(['getent', 'hosts', host],
                            capture_output=True, text=True, timeout=5)
            if getent.returncode == 0:
                test["resolved_ip"] = getent.stdout.strip().split()[0]
        except:
            pass
        # Ollama API test
        try:
            async with aiohttp.ClientSession() as session:
                async with session.get(f"http://{host}:{ollama_port}/api/version",
                                       timeout=aiohttp.ClientTimeout(total=3)) as resp:
                    if resp.status == 200:
                        test["ollama"] = True
                        data = await resp.json()
                        test["ollama_version"] = data.get("version", "unknown")
        except:
            pass
        results["tests"].append(test)

    return results

@app.get("/api/jobs/{job_id}/events")
async def api_sse(request: Request, job_id: str):
    if job_id not in jobs:
        raise HTTPException(404)
    q = asyncio.Queue(maxsize=200)
    sse_queues.setdefault(job_id, []).append(q)
    async def gen():
        try:
            yield f"data: {json.dumps(jobs[job_id].to_dict())}\n\n"
            while True:
                if await request.is_disconnected():
                    break
                try:
                    data = await asyncio.wait_for(q.get(), timeout=15)
                    yield f"data: {json.dumps(data)}\n\n"
                    if data.get("status") in ("completed", "failed", "cancelled"):
                        break
                except asyncio.TimeoutError:
                    yield ": keepalive\n\n"
        finally:
            if job_id in sse_queues and q in sse_queues[job_id]:
                sse_queues[job_id].remove(q)
    return StreamingResponse(gen(), media_type="text/event-stream",
                             headers={"Cache-Control": "no-cache", "X-Accel-Buffering": "no"})

# ============================================================
# Scheduled Cleanup (daily at midnight)
# ============================================================

def cleanup_old_files():
    """Delete all files from uploads/ and outputs/ directories."""
    deleted = 0
    for d in [UPLOAD_DIR, OUTPUT_DIR]:
        if d.exists():
            for f in d.iterdir():
                if f.is_file():
                    try:
                        f.unlink()
                        deleted += 1
                    except Exception as e:
                        app_logger.error(f"Cleanup failed for {f}: {e}")
    app_logger.info(f"Scheduled cleanup: {deleted} files deleted")
    return deleted

def schedule_daily_cleanup():
    """Schedule cleanup to run daily at midnight."""
    def run_cleanup():
        while True:
            now = datetime.now()
            midnight = (now + timedelta(days=1)).replace(
                hour=0, minute=0, second=0, microsecond=0
            )
            wait_seconds = (midnight - now).total_seconds()
            app_logger.info(
                f"Next cleanup scheduled in {wait_seconds:.0f}s (at midnight)"
            )
            time.sleep(wait_seconds)
            cleanup_old_files()

    t = threading.Thread(target=run_cleanup, daemon=True)
    t.start()
    app_logger.info("Daily cleanup scheduler started (runs at midnight)")

# ============================================================
# Main
# ============================================================

if __name__ == "__main__":
    import uvicorn
    port = int(os.environ.get("PORT", 8080))
    schedule_daily_cleanup()
    print("=" * 60)
    print("  Offline Document Translator - Server")
    print(f"  Web UI:       http://0.0.0.0:{port}")
    print(f"  Ollama:       {DEFAULT_OLLAMA_URL} (external)")
    print(f"  LibreTranslate: {DEFAULT_LIBRETRANSLATE_URL} (internal)")
    print(f"  Cleanup:      daily at midnight")
    print("=" * 60)
    uvicorn.run(app, host="0.0.0.0", port=port, log_level="info")
